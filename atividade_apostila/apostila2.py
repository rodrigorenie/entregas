import string
import docx
import nltk

from nltk.corpus import machado
from nltk import PorterStemmer, LancasterStemmer, RSLPStemmer
from nltk.corpus import stopwords


def ex01():
    # Rodrigo Renie de Braga Pinto
    # TEXT ANALYSIS(Apostila)Parte 2.docx
    # Exercitando 1
    # Execute o que se pede.
    # Utilize o arquivo com a obra Memórias Póstumas de Bras Cubas e compare os
    # stemms obtidos a partir do 6º e 7º parágrafos do texto. Utilize Porter,
    # Lancaster e RSPL.

    #
    # Carrega os parágrafos de Memórias Póstumas de Brás Cubas
    #

    book_fileid = 'romance/marm05.txt'

    # Carrega todos os parágrafos do livro, já tokenizados
    book_paras = machado.paras(book_fileid)

    # A posição 17 é o primeiro parágrafo do primeiro capítulo, portanto:
    book_tokens = book_paras[17][0] + book_paras[18][0]

    book_stopwords = stopwords.words('portuguese')
    book_stopwords += [p for p in string.punctuation]
    book_tokens = [t.lower() for t in book_tokens
                   if t.lower() not in book_stopwords]

    #
    # Executa os Stemmers
    #
    porter = PorterStemmer()
    stems_porter = [porter.stem(t) for t in book_tokens]

    lancaster = LancasterStemmer()
    stems_lancaster = [lancaster.stem(t) for t in book_tokens]

    rslp = RSLPStemmer()
    stems_rslp = [rslp.stem(t) for t in book_tokens]

    print('{:18s} {:18s} {:18s} {}'.format('Tokens', 'Porter', 'Lancaster', 'RSLP'))
    for t, p, l, r in zip(book_tokens, stems_porter, stems_lancaster, stems_rslp):
        print('{:18s} {:18s} {:18s} {}'.format(t, p, l, r))


def ex02():
    # Rodrigo Renie de Braga Pinto
    # TEXT ANALYSIS(Apostila)Parte 2.docx
    # Exercitando 2
    # Execute o que se pede.
    # Utilize o arquivo Noticia_1 disponível na pasta de dados da turma e realize
    # as tarefas de POS e NER.
    # Salve o resultado de POS em um arquivo, com <seu_nome>_POS_Noticia1
    # Salve o resultado de NER em um arquivo, com <seu_nome>_NER_Noticia1

    noticia1 = docx.Document('dados/Noticia_1.docx')

    noticia1_tokens = '\n'.join([p.text for p in noticia1.paragraphs])
    noticia1_tokens = nltk.sent_tokenize(noticia1_tokens)
    noticia1_tokens = [nltk.word_tokenize(t) for t in noticia1_tokens]
    print(noticia1_tokens)

    noticia1_pos = [nltk.pos_tag(s) for s in noticia1_tokens]
    noticia1_ner = [nltk.ne_chunk(s) for s in noticia1_pos]
    print(noticia1_pos)

    posdoc = docx.Document()
    for pos in noticia1_pos:
        posdoc.add_paragraph(str(pos))
    posdoc.save('dados/RodrigoRenieDeBragaPinto_POS_Noticia1.docx')

    nerdoc = docx.Document()
    for ner in noticia1_ner:
        nerdoc.add_paragraph(str(ner))
    nerdoc.save('dados/RodrigoRenieDeBragaPinto_NER_Noticia1.docx')


def ex03():
    # Rodrigo Renie de Braga Pinto
    # TEXT ANALYSIS(Apostila)Parte 2.docx
    # Exercitando 3
    # Execute o que se pede.

    topn = 5

    def tree_get_ner(t: nltk.Tree) -> list:
        ner = []
        label = t.label()

        if label != 'S':
            ner = [(' '.join(l for (l, _) in t.leaves()), label)]

        for branch in t:
            if type(branch) == nltk.Tree:
                ner += tree_get_ner(branch)

        return ner

    def tree_get_pos(t: nltk.Tree) -> list:
        return [leave for leave in t.leaves()]

    # 1. Execute print(machado.readme()) para conhecer melhor o corpus
    print(machado.readme())

    # 2. Utilizando oo documentos relativos a Dom Casmurro e O alienista,
    # faça o que se pede
    # Dom Casmurro (1899): romance/marm08.txt
    # Papéis Avultos (1882) (que contém O Alienista): contos/macn003.txt

    # a. Classifique as palavras de acordo com suas classes gramaticais de cada
    # documento. Salve o corpus POS Tagged em uma planilha ou texto para uso
    # posterior. É importante manter a informação sobre o documento origem dos
    # novos documentos;

    print('Gerando POS de livro1')
    livro1 = machado.sents('romance/marm08.txt')
    livro1_pos = [nltk.pos_tag(s) for s in livro1]

    print('Gerando POS de livro2')
    livro2 = machado.sents('contos/macn003.txt')
    livro2_pos = [nltk.pos_tag(s) for s in livro2]

    # b. Obtenha a lista de entidades em cada documento, salvando para uso
    # posterior;

    print('Gerando NER de livro1')
    livro1_forest = [nltk.ne_chunk(s) for s in livro1_pos]

    print('Gerando NER de livro2')
    livro2_forest = [nltk.ne_chunk(s) for s in livro2_pos]

    # c. Analisando os documentos marcados (tagged) tanto POS quanto NER, quais são
    # as classes mais utilizadas?
    livro1_ner = []
    livro1_pos = []
    livro2_ner = []
    livro2_pos = []

    print('Carregando Classes NER e POS de livro1')
    for tree in livro1_forest:
        livro1_ner += tree_get_ner(tree)
        livro1_pos += tree_get_pos(tree)

    print('Carregando Classes NER e POS de livro2')
    for tree in livro2_forest:
        livro2_ner += tree_get_ner(tree)
        livro2_pos += tree_get_pos(tree)

    print('Carregando Frequência das Classes NER e POS de livro1')
    livro1_ner_freq = nltk.FreqDist([t for (_, t) in livro1_ner])
    livro1_pos_freq = nltk.FreqDist([t for (_, t) in livro1_pos])

    print('Carregando Frequência das Classes NER e POS de livro2')
    livro2_ner_freq = nltk.FreqDist([t for (_, t) in livro1_ner])
    livro2_pos_freq = nltk.FreqDist([t for (_, t) in livro1_pos])

    print('\n\n** Frequência das Classes NER de livro1')
    for (a, b) in livro1_ner_freq.most_common(topn):
        print('{:>15} : {:<4}'.format(str(a), b))

    print('\n\n** Frequência das Classes POS de livro1')
    for (a, b) in livro1_pos_freq.most_common(topn):
        print('{:>15} : {:<4}'.format(str(a), b))

    print('\n\n** Frequência das Classes NER de livro2')
    for (a, b) in livro2_ner_freq.most_common(topn):
        print('{:>15} : {:<4}'.format(str(a), b))

    print('\n\n** Frequência das Classes POS de livro2')
    for (a, b) in livro2_pos_freq.most_common(topn):
        print('{:>15} : {:<4}'.format(str(a), b))

    # d. Observe que há uma tendência de que termos menos relevantes para a análise
    # sejam mais frequentes. Então, repita os procedimentos anteriores, mas com os
    # termos que são relevantes para uma análise do que está sendo falado (trata-se
    # de uma análise preliminar e ainda superficial do discurso).

    sw = [p for p in string.punctuation]
    sw += ["''"]

    print('\n\n** Carregando Frequência das Classes NER e POS de livro1 FILTRADO')
    livro1_ner_freq = nltk.FreqDist([t for (_, t) in livro1_ner if t not in sw])
    livro1_pos_freq = nltk.FreqDist([t for (_, t) in livro1_pos if t not in sw])

    print('\n\n** Carregando Frequência das Classes NER e POS de livro2 FILTRADO')
    livro2_ner_freq = nltk.FreqDist([t for (_, t) in livro2_ner if t not in sw])
    livro2_pos_freq = nltk.FreqDist([t for (_, t) in livro2_pos if t not in sw])

    print('\n\n** Frequência das Classes NER de livro1 FILTRADO')
    for (a, b) in livro1_ner_freq.most_common(topn):
        print('{:>15} : {:<4}'.format(str(a), b))

    print('\n\n** Frequência das Classes POS de livro1 FILTRADO')
    for (a, b) in livro1_pos_freq.most_common(topn):
        print('{:>15} : {:<4}'.format(str(a), b))

    print('\n\n** Frequência das Classes NER de livro2 FILTRADO')
    for (a, b) in livro2_ner_freq.most_common(topn):
        print('{:>15} : {:<4}'.format(str(a), b))

    print('\n\n** Frequência das Classes POS de livro2 FILTRADO')
    for (a, b) in livro2_pos_freq.most_common(topn):
        print('{:>15} : {:<4}'.format(str(a), b))

    # e. Determine o vocabulário comum entre os textos

    print('\n\n** Gerando vocabulário comum entre livro1 e livro2')

    # livro1_palavras = [p.lower() for s in livro1 for p in s]
    livro1_palavras = [p.lower() for s in livro1 for p in s if p not in sw]
    livro1_vocab = set(livro1_palavras)

    # livro2_palavras = [p.lower() for s in livro2 for p in s]
    livro2_palavras = [p.lower() for s in livro2 for p in s if p not in sw]
    livro2_vocab = set(livro2_palavras)

    vocab_comum = livro1_vocab & livro2_vocab
    print(vocab_comum)

    # f. Determine a frequência de termos comuns nos dois textos, separadamente

    livro1_palavras_freq = nltk.FreqDist(livro1_palavras)
    livro2_palavras_freq = nltk.FreqDist(livro2_palavras)

    print('\n\n** Frequência das palavras do livro1')
    for (a, b) in livro1_palavras_freq.most_common(topn):
        print('{:>15} : {:<4}'.format(str(a), b))

    print('\n\n** Frequência das palavras do livro2')
    for (a, b) in livro2_palavras_freq.most_common(topn):
        print('{:>15} : {:<4}'.format(str(a), b))

    # g. Determine a frequência de termos comuns utilizados pelo autor considerando
    # os dois textos

    print('\n\n** Gerando palavras em comum entre livro1 e livro2')
    palavras_comum = [p for p in livro1_palavras if p in livro2_palavras]
    palavras_comum += [p for p in livro2_palavras if p in livro1_palavras]
    palavras_comum_freq = nltk.FreqDist(palavras_comum)

    print('\n\n** Frequência das palavras em comum entre livro1 e livro2')
    for (a, b) in palavras_comum_freq.most_common(topn):
        print('{:>15} : {:<4}'.format(str(a), b))

    # h. Desafio: Quais são as entidades mais citadas pelo autor?

    print('\n\n** Gerando entidades em comum entre livro1 e livro2')
    ner_comum_freq = livro1_ner_freq + livro2_ner_freq
    pos_comum_freq = livro1_pos_freq + livro2_pos_freq

    print('\n\n** Frequência das entidades NER em comum entre livro1 e livro2')
    for (a, b) in ner_comum_freq.most_common(topn):
        print('{:>15} : {:<4}'.format(str(a), b))

    print('\n\n** Frequência das entidades POS em comum entre livro1 e livro2')
    for (a, b) in pos_comum_freq.most_common(topn):
        print('{:>15} : {:<4}'.format(str(a), b))
