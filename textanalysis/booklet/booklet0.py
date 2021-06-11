from __future__ import annotations

import docx
import dsutils

from typing import Optional, Union


class Ex01:
    """Implementa a atividade descrita em :ref:`Apostila 0 Exercitando 01`

    :param text: Texto a ser utilizado como base para o exercício.
    :type text: str
    """

    def __init__(self, text: Optional[str] = None) -> None:
        super().__init__()
        self.text = text

    def __str__(self) -> str:
        """Define que a representação de string desta classe é o conteúdo do
        texto
        """
        return self.text

    @property
    def text(self) -> str:
        """Texto a ser utilizado pela classe

        Propriedade ser utilizado pelos outros métodos da classe. Se for
        atribuído ``None``, será convertido para uma string vazia.

        :raises TypeError: Se for atribuído com um valor que não seja string

        :return: Texto a ser utilizado pela classe
        """
        return self._text

    @text.setter
    def text(self, text: str) -> None:
        if text is None:
            text = ''

        if not isinstance(text, str):
            raise TypeError("'text' deve ser uma string")

        self._text = text

    @property
    def text_chars(self) -> iter[str]:
        """Cria um gerador para os caracteres individuais do texto.

        :return: Iterador dos caracteres de :attr:`text`
        """
        for char in self.text:
            yield char

    @property
    def text_split(self) -> iter[str]:
        """Divide o texto em uma lista de palavras

        Divide o texto em uma lista de palavras, separadas por um espaço em
        branco, e cria um gerador para os itens da lista.

        :return: Iterador das palavras de :attr:`text`
        """
        for item in self.text.split():
            yield item

    @property
    def text_split_len(self) -> int:
        """Contabiliza o tamanho da lista do iterador :attr:`text_split`

        :return: Tamanho da lista :attr:`text_split`
        """
        return len(list(self.text_split))

    def text_replace(self, old: str, new: str) -> str:
        """Substitui ``old`` por ``new``

        Substitui o trecho de :attr:`text` indicado pelo parâmetro ``old`` pelo
        texto indicado no parâmetro ``new``

        :param old: Texto original a ser substituído.
        :type old: str
        :param new: Texto novo.
        :type new: str

        :return: O novo valor de :attr:`text`
        """
        self.text = self.text.replace(old, new)
        return self.text

    def text_segment(self, first: int, last: Optional[int] = None) -> str:
        """Retorna o segmento do texto entre ``first`` e ``last``

        Retorna o segmento do texto indicado, da posição inicial ``first``
        até a posição ``last``, ambos *INCLUSIVO*, ou seja, retorna o caracter
        das posições indicadas. Se ``last`` for omitido, retorna apenas o
        caracter indicado por ``first``

        :raises ValueError: Se ``first`` é menor ou igual a zero ou menor que
            ``last``

        :param first: Posição inicial do segmento (deve ser maior que zero)
        :type first: int
        :param last: Posição final do segmento, opcional
        :type last: int, opcional

        :return: String com o segmento do texto ou uma string vazia se
            ``first`` for maior que o tamanho do texto
        """
        if first <= 0:
            raise ValueError("'first' deve ser maior ou igual a 1")

        if not last:
            last = first

        if last < first:
            raise ValueError("'last' deve ser maior ou igual a 'first'")

        return self.text[first-1:last]

    def text_last(self, n: int) -> str:
        """Retorna os últimos caracteres do texto, de tamanho indicado pelo
        parâmetro ``n``, que é opcional com valor padrão ``15``.

        :param n: Tamanho do segmento.
        :type n: int, opcional

        :raises IndexError: se o tamanho do segmento é maior que o texto em si.

        :return: Segmento do texto dos últimos ``n`` caracteres.
        """
        return self.text[-n:]

    def text_save(self, filename: str) -> str:
        """Salva o texto em um arquivo indicado pelo parâmetro ``filename``.

        :param filename: Caminho do arquivo
        :type filename: str

        :return: Caminho completo do arquivo salvo
        """
        filename = dsutils.datadir.join(filename)
        with open(filename, 'w', encoding='utf8') as f:
            f.write(self.text + '\n')
            return filename


class Ex02:
    """Implementa a atividade descrita em :ref:`Apostila 0 Exercitando 02`

    :param docname: Caminho do arquivo ``docx`` a ser carregado.
    :type docname: str
    """

    def __init__(self, docname: str) -> None:
        super().__init__()
        self._doc = docx.Document(dsutils.datadir.join(docname))

    @property
    def paragraphs(self) -> iter[str]:
        """Cria um iterador para os parágrafos encontrados no documento.

        :returns: iterador dos parágrafos no documento
        """

        for paragraph in self._doc.paragraphs:
            yield paragraph.text

    @property
    def paragraphs_list(self) -> list[str]:
        """Cria uma lista contendo cada parágrafo encontrado no documento.

        :return: Lista com os parágrafos.
        """
        return list(self.paragraphs)

    @property
    def paragraphs_len(self) -> int:
        """Contabiliza o tamanho da lista do gerador criado por
        :func:`Exercitando02.paragraphs`.

        :returns: Tamanho da lista.
        """
        return len(self.paragraphs_list)

    def paragraphs_segment(self, first: int,
                           last: Optional[int] = None) -> Union[list[str], str]:
        """Retorna os parágrafos do documento, da posição inicial ``first`` até
        a posição ``last``, ambos *INCLUSIVO*, ou seja, também retorna os
        parágrafos nas posições indicadas. Se ``last`` for omitido, retorna
        apenas o parágrafo indicado por ``first``.

        :raises ValueError: Erro gerado quando ``first`` é menor ou igual a
            zero ou quando ``first`` menor que ``last``.

        :param first: Posição inicial do segmento (deve ser maior que zero).
        :type first: int
        :param last: Posição final do segmento.
        :type last: int, opcional

        :return: Uma lista de string com os parágrafos solicitados, ou uma
            única string de parágrafo de ``last`` for omitido.
        """
        if first <= 0:
            raise ValueError('first parameter must be greater '
                             'or equal to one')

        if last and first > last:
            raise ValueError('last parameter must be greater '
                             'or equal to first')

        if last:
            return self.paragraphs_list[first - 1:last]
        else:
            return self.paragraphs_list[first - 1]

    def paragraphs_hastext(self, text: str) -> bool:
        """Verifica se a string indicado pelo parâmetro ``text`` existe no
        documento.

        :param text: Texto a procurar no documento.
        :type text: str

        :return: Verdadeiro ou falso.
        """
        return text in self.paragraphs_text

    @property
    def paragraphs_text(self) -> str:
        """Retorna uma string com o conteúdo do documento.

        :return: String do documento.
        """
        return '\n'.join(self.paragraphs_list)

    def paragraphs_replacetext(self, old: str, new: str) -> str:
        """Retorna uma string com o conteúdo do documento, substiuindo o texto
        indicado pelo parâmetro ``old`` por ``new``.

        :param old: Texto original a ser substituído.
        :type old: str
        :param new: Texto novo.
        :type new: str

        :return: String com o texto substituído.
        """
        return self.paragraphs_text.replace(old, new)
